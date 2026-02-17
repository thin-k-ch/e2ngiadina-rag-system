#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF → One-Page A4 Summary via Ollama (Map-Reduce)
==================================================

Pipeline:
  1. PDF-Text extrahieren (PyMuPDF → pypdf → pdfplumber Fallback)
  2. Dokumenttyp klassifizieren (technical / contractual)
  3. Text in Chunks aufteilen
  4. MAP: Pro Chunk strukturierte Fakten extrahieren (JSON)
  5. REDUCE: Alle Fakten zu einem One-Pager verdichten

Usage:
  python pdf_onepager.py --pdf /path/to/doc.pdf
  python pdf_onepager.py --pdf /path/to/doc.pdf --model llama4:latest --audience management
  python pdf_onepager.py --pdf /path/to/doc.pdf --out summary.md --format markdown
"""

import argparse
import json
import os
import re
import sys
import time
from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional, Tuple

import requests

# ---------------------------------------------------------------------------
# PDF Extraction (3-tier fallback)
# ---------------------------------------------------------------------------

def extract_pdf_text(pdf_path: str) -> Tuple[str, str]:
    """
    Best-effort PDF text extraction with 3 fallbacks.
    Returns (text, extractor_name).
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF nicht gefunden: {pdf_path}")

    # Tier 1: PyMuPDF (best quality on most PDFs)
    try:
        import fitz
        doc = fitz.open(pdf_path)
        parts = [page.get_text("text") for page in doc]
        text = _normalize_text("\n".join(parts))
        if len(text.strip()) > 200:
            return text, "PyMuPDF"
    except ImportError:
        pass
    except Exception:
        pass

    # Tier 2: pdfplumber (good for tables/structured PDFs)
    try:
        import pdfplumber
        with pdfplumber.open(pdf_path) as pdf:
            parts = [p.extract_text() or "" for p in pdf.pages]
        text = _normalize_text("\n".join(parts))
        if len(text.strip()) > 200:
            return text, "pdfplumber"
    except ImportError:
        pass
    except Exception:
        pass

    # Tier 3: pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(pdf_path)
        parts = [p.extract_text() or "" for p in reader.pages]
        text = _normalize_text("\n".join(parts))
        if len(text.strip()) > 200:
            return text, "pypdf"
    except ImportError:
        pass
    except Exception as e:
        raise RuntimeError(
            "Konnte PDF nicht lesen. Installiere: pip install pymupdf pdfplumber pypdf"
        ) from e

    raise RuntimeError(
        f"PDF enthält zu wenig extrahierbaren Text (evtl. reiner Scan). "
        f"Datei: {pdf_path}"
    )


def _normalize_text(text: str) -> str:
    """Clean up extracted PDF text."""
    text = text.replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove standalone page numbers and very short artifacts
    lines = []
    for ln in text.split("\n"):
        stripped = ln.strip()
        if re.fullmatch(r"\d{1,4}", stripped):
            continue
        if re.fullmatch(r"[-–—_]{3,}", stripped):
            continue
        lines.append(stripped)
    return "\n".join(lines).strip()


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------

def approx_tokens(s: str) -> int:
    """Rough heuristic: ~3.5 chars/token for German text."""
    return max(1, int(len(s) / 3.5))


def split_into_chunks(
    text: str, target_tokens: int = 1800, overlap_tokens: int = 200
) -> List[str]:
    """Split by paragraphs, pack into chunks ≈ target_tokens with overlap."""
    # Try double-newline split first, fall back to single-newline for pdfplumber output
    paras = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if len(paras) <= 2 and len(text) > 3000:
        # pdfplumber often produces single-newline text → split by lines, group ~10 lines
        lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
        paras = []
        for i in range(0, len(lines), 10):
            paras.append("\n".join(lines[i:i+10]))
    chunks: List[str] = []
    cur: List[str] = []
    cur_toks = 0

    for p in paras:
        pt = approx_tokens(p)
        if cur_toks + pt > target_tokens and cur:
            chunks.append("\n\n".join(cur))
            cur = []
            cur_toks = 0
        cur.append(p)
        cur_toks += pt

    if cur:
        chunks.append("\n\n".join(cur))

    # Overlap: prepend tail of previous chunk to each subsequent chunk
    if overlap_tokens > 0 and len(chunks) > 1:
        overlapped = [chunks[0]]
        for i in range(1, len(chunks)):
            tail_chars = overlap_tokens * 4
            tail = chunks[i - 1][-tail_chars:]
            overlapped.append(f"{tail}\n\n{chunks[i]}")
        chunks = overlapped

    return chunks


# ---------------------------------------------------------------------------
# Ollama Client (with retry + dynamic num_ctx)
# ---------------------------------------------------------------------------

@dataclass
class OllamaClient:
    base_url: str = "http://localhost:11434"
    default_timeout: float = 300.0
    max_retries: int = 2

    def chat(
        self, model: str, messages: List[Dict[str, str]],
        options: Optional[Dict[str, Any]] = None,
        timeout: Optional[float] = None,
    ) -> str:
        url = f"{self.base_url}/api/chat"

        # Dynamic num_ctx based on input size
        total_chars = sum(len(m.get("content", "")) for m in messages)
        est_tokens = int(total_chars / 3.5)
        num_ctx = max(4096, est_tokens + 8192)  # input + generous output headroom
        num_ctx = min(num_ctx, 131072)

        opts = {"num_ctx": num_ctx}
        if options:
            opts.update(options)

        payload = {
            "model": model,
            "messages": messages,
            "stream": False,
            "options": opts,
        }

        effective_timeout = timeout or self.default_timeout
        last_err = None

        for attempt in range(1, self.max_retries + 1):
            try:
                r = requests.post(url, json=payload, timeout=effective_timeout)
                r.raise_for_status()
                data = r.json()
                return data["message"]["content"]
            except (requests.Timeout, requests.ConnectionError) as e:
                last_err = e
                _log(f"  ⚠️  Versuch {attempt}/{self.max_retries} fehlgeschlagen: {e}")
                if attempt < self.max_retries:
                    time.sleep(2 * attempt)
            except requests.HTTPError as e:
                raise RuntimeError(f"Ollama HTTP-Fehler: {e}\nResponse: {r.text[:500]}") from e

        raise RuntimeError(f"Ollama nicht erreichbar nach {self.max_retries} Versuchen: {last_err}")


# ---------------------------------------------------------------------------
# Logging (stderr, so stdout stays clean for piping)
# ---------------------------------------------------------------------------

def _log(msg: str):
    print(msg, file=sys.stderr, flush=True)


def _extract_json(text: str) -> dict:
    """
    Robustly extract a JSON object from LLM output that may contain
    markdown fences, explanatory text, or other wrapping.
    """
    if not text or not text.strip():
        raise ValueError("Leere Antwort vom Modell")

    # Strategy 1: Direct parse
    try:
        return json.loads(text.strip())
    except json.JSONDecodeError:
        pass

    # Strategy 2: Strip markdown fences (```json ... ``` or ``` ... ```)
    cleaned = re.sub(r"```(?:json)?\s*\n?", "", text)
    cleaned = re.sub(r"\n?```\s*$", "", cleaned).strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass

    # Strategy 3: Find first { ... last } (greedy)
    first_brace = text.find("{")
    last_brace = text.rfind("}")
    if first_brace != -1 and last_brace > first_brace:
        candidate = text[first_brace:last_brace + 1]
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            pass

    # Strategy 4: Try to fix common issues (trailing commas, single quotes)
    if first_brace != -1 and last_brace > first_brace:
        candidate = text[first_brace:last_brace + 1]
        # Remove trailing commas before } or ]
        candidate = re.sub(r",\s*([}\]])", r"\1", candidate)
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            pass

    raise ValueError(f"Kein valides JSON gefunden in Antwort ({len(text)} Zeichen)")


# ---------------------------------------------------------------------------
# Document Classification
# ---------------------------------------------------------------------------

DOC_CLASSIFIER_SYSTEM = """Du bist ein präziser Dokumenten-Klassifizierer für Schweizer Bau- und Infrastrukturprojekte.
Gib ausschliesslich JSON zurück. Kein Fliesstext, keine Markdown-Fences."""

DOC_CLASSIFIER_USER = """Klassifiziere das Dokument als:
- "technical" (technisch/engineering/IT/Architektur/Spezifikation/Pflichtenheft/Testprotokoll/Incident)
- "contractual" (vertraglich/rechtlich/Werkvertrag/SIA/SLA/AGB/Pönale/Datenschutz/Offerte)

Gib JSON zurück:
{{
  "type": "technical" oder "contractual",
  "confidence": 0.0-1.0,
  "signals": ["max 5 Stichworte die zur Klassifikation führten"]
}}

TEXTAUSZUG (erste ~3000 Wörter):
---
{excerpt}
---"""

# CH-specific keywords for fallback classification
_CONTRACTUAL_KW = {
    "vertrag", "werkvertrag", "agb", "haftung", "vertraulich", "sla", "datenschutz",
    "pönale", "konventionalstrafe", "sia", "vergütung", "abnahme", "gewährleistung",
    "kündigung", "schadenersatz", "pflichtenheft", "offerte", "nachtrag", "mahnung",
    "verzug", "frist", "leistungsverzeichnis",
}


def classify_document(
    client: OllamaClient, model: str, text: str
) -> Tuple[str, float, List[str]]:
    """Classify document as technical or contractual."""
    excerpt = text[:12000]

    _log("🔍 Klassifiziere Dokument...")
    t0 = time.time()

    try:
        content = client.chat(
            model=model,
            messages=[
                {"role": "system", "content": DOC_CLASSIFIER_SYSTEM},
                {"role": "user", "content": DOC_CLASSIFIER_USER.format(excerpt=excerpt)},
            ],
            options={"temperature": 0.0},
            timeout=120.0,
        )
        obj = _extract_json(content)
        doc_type = obj.get("type", "technical")
        conf = float(obj.get("confidence", 0.5))
        signals = [str(s) for s in obj.get("signals", [])][:5]
        if doc_type not in ("technical", "contractual"):
            doc_type = "technical"
    except Exception as e:
        _log(f"  ⚠️  LLM-Klassifikation fehlgeschlagen ({e}), nutze Heuristik")
        lower = excerpt.lower()
        hits = [kw for kw in _CONTRACTUAL_KW if kw in lower]
        if len(hits) >= 2:
            doc_type, conf, signals = "contractual", 0.6, hits[:5]
        else:
            doc_type, conf, signals = "technical", 0.5, ["heuristic:default"]

    _log(f"   → {doc_type} (confidence={conf:.2f}, {time.time()-t0:.1f}s)")
    _log(f"   Signals: {', '.join(signals)}")
    return doc_type, conf, signals


# ---------------------------------------------------------------------------
# MAP Phase: Extract structured facts per chunk
# ---------------------------------------------------------------------------

MAP_SYSTEM = """Du bist ein exakter Assistent für Dokumenten-Analyse in Schweizer Infrastrukturprojekten.
Halte dich strikt an die JSON-Ausgabestruktur.
Keine Halluzinationen: Wenn etwas nicht im Text steht, lass das Feld leer oder schreibe "Nicht erwähnt".
Gib NUR valides JSON zurück, keine Markdown-Fences."""

MAP_USER = """Rolle: {audience}
Dokumenttyp: {doc_type}
Chunk {chunk_idx}/{total_chunks}

Extrahiere aus dem folgenden TEXTCHUNK die wichtigsten Fakten als kompaktes JSON:
{{
  "key_points": ["max 5 Kernaussagen"],
  "entities": ["Firmen/Personen/Systeme/Produkte"],
  "dates_numbers": ["Daten, Beträge, SLAs, KPIs, Fristen, Versionen"],
  "risks_issues": ["Risiken/Offene Punkte/Abhängigkeiten"],
  "actions": ["Konkrete Next Steps/To-dos"],
  "coverage": "1 Satz: welche Themen dieser Chunk abdeckt"
}}

TEXTCHUNK:
---
{chunk}
---"""


def map_extract_facts(
    client: OllamaClient, model: str, doc_type: str,
    audience: str, chunks: List[str],
) -> List[Dict[str, Any]]:
    """MAP phase: extract structured facts from each chunk."""
    facts = []
    total = len(chunks)

    for i, ch in enumerate(chunks, start=1):
        _log(f"  📄 MAP Chunk {i}/{total} ({approx_tokens(ch)} tokens)...")
        t0 = time.time()

        resp = client.chat(
            model=model,
            messages=[
                {"role": "system", "content": MAP_SYSTEM},
                {"role": "user", "content": MAP_USER.format(
                    audience=audience, doc_type=doc_type,
                    chunk=ch[:10000], chunk_idx=i, total_chunks=total,
                )},
            ],
            options={"temperature": 0.1},
        )

        try:
            obj = _extract_json(resp)
            obj["_chunk"] = i
        except Exception as e:
            _log(f"     ⚠️  JSON-Parse fehlgeschlagen ({e}), speichere Rohtext")
            _log(f"     Antwort (erste 200 Zeichen): {resp[:200]}")
            obj = {
                "_chunk": i,
                "key_points": [resp[:500] if resp else "Parsing-Fehler"],
                "entities": [], "dates_numbers": [],
                "risks_issues": [], "actions": [],
                "coverage": "Parse-Fehler",
            }

        facts.append(obj)
        _log(f"     ✅ {len(obj.get('key_points', []))} Punkte ({time.time()-t0:.1f}s)")

    return facts


# ---------------------------------------------------------------------------
# REDUCE Phase: Synthesize one-pager from all facts
# ---------------------------------------------------------------------------

REDUCE_SYSTEM = """Du bist ein Senior Technical Writer für Schweizer Infrastrukturprojekte.
Erzeuge eine einseitige A4-Zusammenfassung (deutsch).
Regeln:
- Ziel: 450–650 Wörter (max. 1 A4-Seite)
- Prägnant, klar, strukturiert
- KEINE erfundenen Details, KEINE Vermutungen, KEINE generischen Empfehlungen
- Wenn etwas nicht explizit im Dokument steht, schreibe "Nicht im Dokument erwähnt"
- Erfinde KEINE Handlungsempfehlungen oder Nächsten Schritte die nicht im Dokument stehen
- Verwende Fachbegriffe korrekt (SIA, FAT, SAT, TFK, etc.)"""

REDUCE_USER = """Rolle: {audience}
Dokumenttyp: {doc_type}
Quelle: {filename}

Verdichte die folgenden extrahierten Chunk-Fakten zu EINEM One-Pager.

Ausgabeformat (genau so, mit Überschriften):

{format_template}

Chunk-Fakten ({n_chunks} Abschnitte):
{facts_json}"""

FORMAT_MARKDOWN_TECHNICAL = """# [Titel des Dokuments]

## Kontext & Zweck
2–3 Sätze

## Kernaussagen
- Bullet 1
- ...  (5–8 Bullets)

## Wichtige Zahlen / Daten / Fristen
- Bullet (falls vorhanden; sonst "Nicht im Dokument erwähnt")

## Risiken & offene Punkte
- 3–6 Bullets (NUR was im Dokument steht; sonst "Keine explizit genannt")

## Pendenzen / Nächste Schritte
- NUR wenn im Dokument explizit erwähnt; sonst "Keine im Dokument definiert"

## Relevante Entitäten
Kompakt, komma-separiert"""

FORMAT_MARKDOWN_CONTRACTUAL = """# [Titel des Dokuments]

## Kontext & Zweck
2–3 Sätze

## Kernaussagen
- Bullet 1
- ...  (5–8 Bullets)

## Wichtige Zahlen / Daten / Fristen
- Bullet (falls vorhanden; sonst "Nicht im Dokument erwähnt")

## Vertragliche Pflichten & Konditionen
- Wesentliche Pflichten beider Parteien (NUR was im Vertrag steht)

## Risiken & offene Punkte
- 3–6 Bullets (NUR was im Dokument steht; sonst "Keine explizit genannt")

## Pendenzen
- NUR wenn im Dokument explizit erwähnt; sonst "Keine im Dokument definiert"

## Relevante Entitäten
Kompakt, komma-separiert"""

FORMAT_PLAIN = """1) Titel (falls nicht erkennbar: "Dokument-Übersicht")
2) Kontext & Zweck (2–3 Sätze)
3) Kernaussagen (5–8 Bullet Points)
4) Wichtige Zahlen/Daten/Fristen (Bullets, falls vorhanden; sonst "Nicht im Dokument erwähnt")
5) Risiken & offene Punkte (NUR was im Dokument steht; sonst "Keine explizit genannt")
6) Pendenzen (NUR wenn explizit im Dokument; sonst "Keine im Dokument definiert")
7) Anhang: Relevante Entitäten/Begriffe (kompakt, komma-separiert)"""


def reduce_onepager(
    client: OllamaClient, model: str, doc_type: str,
    audience: str, facts: List[Dict[str, Any]],
    filename: str = "", output_format: str = "markdown",
) -> str:
    """REDUCE phase: synthesize all facts into a one-page summary."""
    facts_json = json.dumps(facts, ensure_ascii=False, indent=1)

    # Truncate facts if extremely long (>30K chars → summarize facts first)
    if len(facts_json) > 30000:
        _log(f"  ⚠️  Fakten sehr lang ({len(facts_json)} chars), kürze...")
        facts_json = facts_json[:30000] + "\n... (gekürzt)"

    if output_format in ("markdown", "docx"):
        fmt = FORMAT_MARKDOWN_CONTRACTUAL if doc_type == "contractual" else FORMAT_MARKDOWN_TECHNICAL
    else:
        fmt = FORMAT_PLAIN

    _log(f"  📝 REDUCE: Erstelle One-Pager ({output_format})...")
    t0 = time.time()

    resp = client.chat(
        model=model,
        messages=[
            {"role": "system", "content": REDUCE_SYSTEM},
            {"role": "user", "content": REDUCE_USER.format(
                audience=audience, doc_type=doc_type,
                filename=filename, n_chunks=len(facts),
                facts_json=facts_json, format_template=fmt,
            )},
        ],
        options={"temperature": 0.2},
    )

    _log(f"     ✅ One-Pager generiert ({time.time()-t0:.1f}s)")
    return resp.strip()


# ---------------------------------------------------------------------------
# Audience selection
# ---------------------------------------------------------------------------

def choose_audience(doc_type: str, audience_arg: str) -> str:
    if audience_arg in ("management", "lead_engineer"):
        return audience_arg
    return "lead_engineer" if doc_type == "technical" else "management"


# ---------------------------------------------------------------------------
# DOCX Output
# ---------------------------------------------------------------------------

def _write_docx(markdown_text: str, out_path: str, meta: dict):
    """Convert markdown one-pager to a formatted Word document."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()

    # Page margins (A4, narrow)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    # Heading styles
    for level, size, color in [(1, 16, "1F4E79"), (2, 12, "2E75B6"), (3, 11, "404040")]:
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = "Calibri"
        h_style.font.size = Pt(size)
        h_style.font.color.rgb = RGBColor.from_string(color)
        h_style.font.bold = True
        h_style.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        h_style.paragraph_format.space_after = Pt(4)

    # Parse markdown and write to docx
    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Empty line → skip
        if not stripped:
            i += 1
            continue

        # Heading
        if stripped.startswith("# ") and not stripped.startswith("## "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        # Bullet point (- or *)
        if re.match(r"^[-*]\s", stripped):
            bullet_text = re.sub(r"^[-*]\s+", "", stripped)
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, bullet_text)
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+[.)]\s", stripped):
            list_text = re.sub(r"^\d+[.)]\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, list_text)
            i += 1
            continue

        # Normal paragraph (may span multiple lines)
        para_lines = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith(("#", "-", "*")):
            if re.match(r"^\d+[.)]\s", lines[i].strip()):
                break
            para_lines.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        _add_formatted_text(p, " ".join(para_lines))
        continue

    # Footer: metadata
    doc.add_paragraph()  # spacer
    footer_p = doc.add_paragraph()
    footer_p.paragraph_format.space_before = Pt(12)
    run = footer_p.add_run(
        f"─── Generiert: {time.strftime('%d.%m.%Y %H:%M')} │ "
        f"Modell: {meta.get('model', '?')} │ "
        f"Typ: {meta.get('doc_type', '?')} ({meta.get('confidence', 0):.0%}) │ "
        f"Zielgruppe: {meta.get('audience', '?')} │ "
        f"Dauer: {meta.get('elapsed', 0):.0f}s ───"
    )
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(150, 150, 150)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out_path)
    _log(f"     📄 DOCX: {os.path.getsize(out_path):,} bytes")


def _add_formatted_text(paragraph, text: str):
    """Add text to a paragraph, handling **bold** and *italic* markdown."""
    # Split on bold (**...**) and italic (*...*)
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

@dataclass
class ProcessResult:
    """Result of processing a single PDF."""
    pdf_path: str
    success: bool
    doc_type: str = ""
    confidence: float = 0.0
    audience: str = ""
    chunks: int = 0
    extractor: str = ""
    elapsed: float = 0.0
    output_path: str = ""
    error: str = ""


def process_single_pdf(
    pdf_path: str,
    client: OllamaClient,
    model: str,
    audience_arg: str = "auto",
    output_format: str = "markdown",
    out_path: str = "",
    chunk_tokens: int = 1800,
    overlap_tokens: int = 200,
) -> ProcessResult:
    """
    Process a single PDF through the full Map-Reduce pipeline.
    Returns a ProcessResult with metadata.
    """
    t_start = time.time()

    try:
        # 1. Extract
        _log(f"📄 Extrahiere PDF-Text...")
        text, extractor = extract_pdf_text(pdf_path)
        n_pages = text.count("\n\n") // 2 or 1
        _log(f"   → {len(text):,} Zeichen via {extractor} (~{n_pages} Seiten)")

        # 2. Classify
        doc_type, conf, signals = classify_document(client, model, text)
        audience = choose_audience(doc_type, audience_arg)
        _log(f"   Audience: {audience}")

        # 3. Chunk
        chunks = split_into_chunks(text, target_tokens=chunk_tokens, overlap_tokens=overlap_tokens)
        _log(f"\n📦 {len(chunks)} Chunks erstellt")

        # 4. MAP
        _log(f"\n🔬 MAP Phase ({len(chunks)} Chunks):")
        facts = map_extract_facts(client, model, doc_type, audience, chunks)

        # 5. REDUCE
        _log(f"\n📝 REDUCE Phase:")
        filename = os.path.basename(pdf_path)
        onepager = reduce_onepager(
            client, model, doc_type, audience, facts,
            filename=filename, output_format=output_format,
        )

        # 6. Output
        elapsed = time.time() - t_start
        meta = {
            "pdf": pdf_path,
            "model": model,
            "doc_type": doc_type,
            "confidence": conf,
            "audience": audience,
            "chunks": len(chunks),
            "extractor": extractor,
            "elapsed": elapsed,
        }

        actual_out = out_path
        if output_format == "docx":
            actual_out = out_path or os.path.splitext(os.path.basename(pdf_path))[0] + "_onepager.docx"
            _write_docx(onepager, actual_out, meta)
            _log(f"\n✅ Word-Dokument gespeichert: {actual_out}")
        else:
            header = (
                f"<!-- PDF One-Pager -->\n"
                f"<!-- PDF: {pdf_path} -->\n"
                f"<!-- Model: {model} | Type: {doc_type} ({conf:.0%}) | Audience: {audience} -->\n"
                f"<!-- Chunks: {len(chunks)} | Extractor: {extractor} | Zeit: {elapsed:.0f}s -->\n\n"
            )
            output = header + onepager + "\n"
            if actual_out:
                with open(actual_out, "w", encoding="utf-8") as f:
                    f.write(output)
                _log(f"\n✅ Gespeichert: {actual_out}")
            else:
                print(output)

        _log(f"\n=== Fertig in {elapsed:.0f}s ({len(chunks)} Chunks, {doc_type}/{audience}) ===")
        return ProcessResult(
            pdf_path=pdf_path, success=True, doc_type=doc_type,
            confidence=conf, audience=audience, chunks=len(chunks),
            extractor=extractor, elapsed=elapsed, output_path=actual_out,
        )

    except Exception as e:
        elapsed = time.time() - t_start
        _log(f"\n❌ Fehler bei {pdf_path}: {e}")
        return ProcessResult(
            pdf_path=pdf_path, success=False, elapsed=elapsed, error=str(e),
        )


def run_batch(
    dir_path: str,
    client: OllamaClient,
    model: str,
    audience_arg: str = "auto",
    output_format: str = "markdown",
    out_dir: str = "",
    chunk_tokens: int = 1800,
    overlap_tokens: int = 200,
    recursive: bool = False,
) -> List[ProcessResult]:
    """
    Process all PDFs in a directory. Returns list of results.
    """
    if not os.path.isdir(dir_path):
        raise FileNotFoundError(f"Verzeichnis nicht gefunden: {dir_path}")

    # Collect PDFs
    pdfs = []
    if recursive:
        for root, _, files in os.walk(dir_path):
            for f in sorted(files):
                if f.lower().endswith(".pdf"):
                    pdfs.append(os.path.join(root, f))
    else:
        pdfs = sorted(
            os.path.join(dir_path, f) for f in os.listdir(dir_path)
            if f.lower().endswith(".pdf")
        )

    if not pdfs:
        _log(f"⚠️  Keine PDFs gefunden in: {dir_path}")
        return []

    # Prepare output directory
    if not out_dir:
        out_dir = os.path.join(dir_path, "_summaries")
    os.makedirs(out_dir, exist_ok=True)

    ext = ".docx" if output_format == "docx" else ".md" if output_format == "markdown" else ".txt"

    _log(f"\n{'='*60}")
    _log(f"📁 BATCH-MODUS: {len(pdfs)} PDFs in {dir_path}")
    _log(f"   Output: {out_dir}")
    _log(f"   Format: {output_format} | Modell: {model}")
    _log(f"{'='*60}\n")

    results: List[ProcessResult] = []
    t_batch_start = time.time()

    for i, pdf_path in enumerate(pdfs, start=1):
        pdf_name = os.path.basename(pdf_path)
        out_name = os.path.splitext(pdf_name)[0] + "_onepager" + ext
        out_path = os.path.join(out_dir, out_name)

        # Skip if already exists
        if os.path.exists(out_path):
            _log(f"\n⏭️  [{i}/{len(pdfs)}] Überspringe (existiert): {pdf_name}")
            results.append(ProcessResult(
                pdf_path=pdf_path, success=True, output_path=out_path,
                doc_type="skipped", elapsed=0,
            ))
            continue

        _log(f"\n{'─'*50}")
        _log(f"📄 [{i}/{len(pdfs)}] {pdf_name}")
        _log(f"{'─'*50}")

        result = process_single_pdf(
            pdf_path=pdf_path, client=client, model=model,
            audience_arg=audience_arg, output_format=output_format,
            out_path=out_path, chunk_tokens=chunk_tokens,
            overlap_tokens=overlap_tokens,
        )
        results.append(result)

    # Summary report
    batch_elapsed = time.time() - t_batch_start
    ok = [r for r in results if r.success]
    skipped = [r for r in results if r.doc_type == "skipped"]
    failed = [r for r in results if not r.success]

    _log(f"\n{'='*60}")
    _log(f"📊 BATCH-ERGEBNIS")
    _log(f"   Gesamt: {len(pdfs)} PDFs | ✅ {len(ok)} OK | ⏭️ {len(skipped)} übersprungen | ❌ {len(failed)} Fehler")
    _log(f"   Dauer: {batch_elapsed:.0f}s ({batch_elapsed/60:.1f} min)")
    _log(f"   Output: {out_dir}")
    if failed:
        _log(f"\n   Fehlgeschlagene PDFs:")
        for r in failed:
            _log(f"     ❌ {os.path.basename(r.pdf_path)}: {r.error}")
    _log(f"{'='*60}")

    # Write combined document (all one-pagers in one file, 1 A4 page per PDF)
    combined_path = os.path.join(out_dir, "_alle_summaries.md")
    with open(combined_path, "w", encoding="utf-8") as f:
        f.write(f"<!-- Batch-Zusammenfassung: {len(pdfs)} Dokumente | {time.strftime('%d.%m.%Y %H:%M')} -->\n\n")
        for r in results:
            if r.success and r.output_path and os.path.exists(r.output_path):
                try:
                    content = open(r.output_path, "r", encoding="utf-8").read()
                    f.write(content.strip() + "\n\n")
                    f.write("---\n<!-- page-break -->\n\n")
                except Exception:
                    pass
    _log(f"   Kombiniert: {combined_path} ({len([r for r in results if r.success])} Seiten)")

    # Write batch report
    report_path = os.path.join(out_dir, "_batch_report.md")
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(f"# Batch-Zusammenfassung\n\n")
        f.write(f"- **Verzeichnis**: {dir_path}\n")
        f.write(f"- **Modell**: {model}\n")
        f.write(f"- **Datum**: {time.strftime('%d.%m.%Y %H:%M')}\n")
        f.write(f"- **Gesamt**: {len(pdfs)} PDFs | ✅ {len(ok)} | ⏭️ {len(skipped)} | ❌ {len(failed)}\n")
        f.write(f"- **Dauer**: {batch_elapsed:.0f}s\n\n")
        f.write(f"| # | PDF | Status | Typ | Chunks | Dauer |\n")
        f.write(f"|---|-----|--------|-----|--------|-------|\n")
        for i, r in enumerate(results, 1):
            name = os.path.basename(r.pdf_path)[:40]
            status = "⏭️" if r.doc_type == "skipped" else ("✅" if r.success else "❌")
            f.write(f"| {i} | {name} | {status} | {r.doc_type} | {r.chunks} | {r.elapsed:.0f}s |\n")
        if failed:
            f.write(f"\n## Fehler\n\n")
            for r in failed:
                f.write(f"- **{os.path.basename(r.pdf_path)}**: {r.error}\n")
    _log(f"   Report: {report_path}")

    return results


def main():
    ap = argparse.ArgumentParser(
        description="PDF → One-Page A4 Summary via Ollama (Map-Reduce)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""Beispiele:
  # Einzelne PDF:
  python pdf_onepager.py --pdf Werkvertrag.pdf
  python pdf_onepager.py --pdf Pflichtenheft.pdf --model llama4:latest --audience lead_engineer
  python pdf_onepager.py --pdf Vertrag.pdf --format docx --out zusammenfassung.docx

  # Ganzer Ordner (Batch):
  python pdf_onepager.py --dir /pfad/zu/pdfs/
  python pdf_onepager.py --dir /pfad/zu/pdfs/ --format docx --out-dir /pfad/zu/output/
  python pdf_onepager.py --dir /pfad/zu/pdfs/ --recursive"""
    )
    group = ap.add_mutually_exclusive_group(required=True)
    group.add_argument("--pdf", help="Pfad zur PDF-Datei (Einzelmodus)")
    group.add_argument("--dir", help="Pfad zum Verzeichnis mit PDFs (Batch-Modus)")
    ap.add_argument("--model", default="llama4:latest", help="Ollama Modell (default: llama4:latest)")
    ap.add_argument("--base-url", default="http://localhost:11434",
                    help="Ollama API URL (default: http://localhost:11434)")
    ap.add_argument("--audience", default="auto", choices=["auto", "management", "lead_engineer"],
                    help="Zielgruppe (default: auto → je nach Dokumenttyp)")
    ap.add_argument("--format", default="markdown", choices=["markdown", "plain", "docx"],
                    help="Ausgabeformat (default: markdown). 'docx' erzeugt ein Word-Dokument.")
    ap.add_argument("--out", default="", help="Output-Datei (nur Einzelmodus). Leer = stdout.")
    ap.add_argument("--out-dir", default="", help="Output-Verzeichnis (nur Batch-Modus). Leer = _summaries im Quellordner.")
    ap.add_argument("--recursive", action="store_true", help="Unterordner einschliessen (nur Batch-Modus)")
    ap.add_argument("--chunk-tokens", type=int, default=1800,
                    help="Ziel-Tokens pro Chunk (default: 1800)")
    ap.add_argument("--overlap-tokens", type=int, default=200,
                    help="Overlap-Tokens zwischen Chunks (default: 200)")
    ap.add_argument("--timeout", type=float, default=300.0,
                    help="Timeout pro LLM-Call in Sekunden (default: 300)")
    args = ap.parse_args()

    client = OllamaClient(base_url=args.base_url, default_timeout=args.timeout)

    if args.dir:
        # Batch mode
        results = run_batch(
            dir_path=args.dir, client=client, model=args.model,
            audience_arg=args.audience, output_format=args.format,
            out_dir=args.out_dir, chunk_tokens=args.chunk_tokens,
            overlap_tokens=args.overlap_tokens, recursive=args.recursive,
        )
        failed = [r for r in results if not r.success]
        if failed:
            sys.exit(1)
    else:
        # Single PDF mode
        _log(f"=== PDF One-Pager ===")
        _log(f"PDF:   {args.pdf}")
        _log(f"Model: {args.model}")
        _log(f"URL:   {args.base_url}")
        _log("")

        result = process_single_pdf(
            pdf_path=args.pdf, client=client, model=args.model,
            audience_arg=args.audience, output_format=args.format,
            out_path=args.out, chunk_tokens=args.chunk_tokens,
            overlap_tokens=args.overlap_tokens,
        )
        if not result.success:
            sys.exit(1)


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        sys.exit(130)
    except Exception as e:
        _log(f"\n❌ Fehler: {e}")
        sys.exit(1)
