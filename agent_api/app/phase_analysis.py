"""
phase_analysis.py

Phase 3: Document Analysis Agent
Reads and analyzes PDF, DOCX, EML files with attachments.
Extracts structured information (FAT/SAT/TIB findings, etc.)
Provides streaming progress updates.
"""

from __future__ import annotations

import os
import re
from typing import Any, Dict, List, AsyncGenerator, Optional
from dataclasses import dataclass


@dataclass
class DocumentSection:
    """Represents a section of a document"""
    page: int
    heading: str
    text: str
    section_type: str  # "paragraph", "table", "header", "footer"
    metadata: Dict[str, Any]


class DocumentNormalizer:
    """Normalizes various document formats to common structure"""
    
    def __init__(self, file_base: str = ""):
        self.file_base = file_base or os.getenv("FILE_BASE", "/media/felix/RAG/1")
    
    def normalize(self, path: str) -> Dict[str, Any]:
        """
        Normalize document to common format:
        {
            "doc_id": "...",
            "type": "pdf|docx|eml",
            "path": "...",
            "sections": [...],
            "metadata": {...},
            "extracted_findings": []
        }
        """
        ext = os.path.splitext(path)[1].lower()
        
        if ext == ".pdf":
            return self._read_pdf(path)
        elif ext in [".docx", ".doc"]:
            return self._read_docx(path)
        elif ext in [".eml", ".msg"]:
            return self._read_eml(path)
        elif ext in [".txt", ".md"]:
            return self._read_text(path)
        else:
            return self._create_empty_document(path, ext)
    
    def _read_pdf(self, path: str) -> Dict[str, Any]:
        """Read PDF with layout-aware extraction"""
        try:
            import fitz  # PyMuPDF
            
            full_path = self._resolve_path(path)
            doc = fitz.open(full_path)
            
            sections = []
            all_text = []
            
            for page_num, page in enumerate(doc, 1):
                # Extract text with layout info
                blocks = page.get_text("dict").get("blocks", [])
                
                page_sections = []
                for block in blocks:
                    if "lines" not in block:
                        continue
                    
                    # Extract text from block
                    text = "\n".join(
                        span["text"] 
                        for line in block["lines"] 
                        for span in line["spans"]
                    )
                    
                    if not text.strip():
                        continue
                    
                    # Detect if this is a heading (large font or bold)
                    is_heading = False
                    for line in block["lines"]:
                        for span in line["spans"]:
                            if span.get("size", 11) > 13 or span.get("flags", 0) & 2**4:
                                is_heading = True
                                break
                    
                    section_type = "heading" if is_heading else "paragraph"
                    
                    page_sections.append(DocumentSection(
                        page=page_num,
                        heading="",
                        text=text.strip(),
                        section_type=section_type,
                        metadata={"bbox": block.get("bbox")}
                    ))
                    
                    all_text.append(text)
                
                sections.extend(page_sections)
            
            doc.close()
            
            return {
                "doc_id": path,
                "type": "pdf",
                "path": path,
                "sections": [
                    {
                        "page": s.page,
                        "heading": s.heading,
                        "text": s.text,
                        "type": s.section_type
                    }
                    for s in sections
                ],
                "metadata": {
                    "page_count": len([s for s in sections if s.page > 0]),
                    "char_count": sum(len(s.text) for s in sections),
                    "full_text": "\n".join(all_text)
                },
                "extracted_findings": []
            }
            
        except Exception as e:
            return self._create_error_document(path, "pdf", str(e))
    
    def _read_docx(self, path: str) -> Dict[str, Any]:
        """Read DOCX with structure extraction"""
        try:
            from docx import Document
            
            full_path = self._resolve_path(path)
            doc = Document(full_path)
            
            sections = []
            all_text = []
            page_num = 1  # DOCX doesn't have explicit pages
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Detect heading by style
                is_heading = para.style.name.startswith("Heading") if para.style else False
                
                sections.append(DocumentSection(
                    page=page_num,
                    heading=para.style.name if is_heading else "",
                    text=text,
                    section_type="heading" if is_heading else "paragraph",
                    metadata={"style": para.style.name if para.style else None}
                ))
                
                all_text.append(text)
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables, 1):
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                
                if table_text:
                    sections.append(DocumentSection(
                        page=page_num,
                        heading=f"Table {table_idx}",
                        text="\n".join(table_text),
                        section_type="table",
                        metadata={"table_index": table_idx}
                    ))
            
            return {
                "doc_id": path,
                "type": "docx",
                "path": path,
                "sections": [
                    {
                        "page": s.page,
                        "heading": s.heading,
                        "text": s.text,
                        "type": s.section_type
                    }
                    for s in sections
                ],
                "metadata": {
                    "paragraph_count": len([s for s in sections if s.section_type == "paragraph"]),
                    "table_count": len([s for s in sections if s.section_type == "table"]),
                    "char_count": sum(len(s.text) for s in sections),
                    "full_text": "\n".join(all_text)
                },
                "extracted_findings": []
            }
            
        except Exception as e:
            return self._create_error_document(path, "docx", str(e))
    
    def _read_eml(self, path: str) -> Dict[str, Any]:
        """Read EML with attachment handling"""
        try:
            import email
            from email import policy
            from email.parser import BytesParser
            
            full_path = self._resolve_path(path)
            
            with open(full_path, 'rb') as f:
                msg = BytesParser(policy=policy.default).parse(f)
            
            # Extract body
            body_text = ""
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    if content_type == "text/plain":
                        body_text = part.get_content()
                        break
                    elif content_type == "text/html":
                        # Simple HTML to text
                        html = part.get_content()
                        body_text = self._html_to_text(html)
                        break
            else:
                body_text = msg.get_content()
            
            # Process attachments
            attachments = []
            normalized_attachments = []
            
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_disposition() == "attachment":
                        filename = part.get_filename()
                        if filename:
                            attachments.append(filename)
                            
                            # Try to normalize attachment if it's a document
                            ext = os.path.splitext(filename)[1].lower()
                            if ext in ['.pdf', '.docx', '.doc', '.txt']:
                                # Save attachment temporarily and process
                                content = part.get_payload(decode=True)
                                if content:
                                    temp_path = f"/tmp/{filename}"
                                    with open(temp_path, 'wb') as f:
                                        f.write(content)
                                    normalized_attachments.append(self.normalize(temp_path))
            
            sections = [
                DocumentSection(
                    page=1,
                    heading=f"Subject: {msg['Subject'] or 'No Subject'}",
                    text=body_text,
                    section_type="header",
                    metadata={"from": msg["From"], "to": msg["To"], "date": msg["Date"]}
                )
            ]
            
            return {
                "doc_id": path,
                "type": "eml",
                "path": path,
                "sections": [
                    {
                        "page": s.page,
                        "heading": s.heading,
                        "text": s.text,
                        "type": s.section_type
                    }
                    for s in sections
                ],
                "metadata": {
                    "subject": msg["Subject"],
                    "from": msg["From"],
                    "to": msg["To"],
                    "date": msg["Date"],
                    "attachments": attachments,
                    "char_count": len(body_text)
                },
                "attachments": normalized_attachments,
                "extracted_findings": []
            }
            
        except Exception as e:
            return self._create_error_document(path, "eml", str(e))
    
    def _read_text(self, path: str) -> Dict[str, Any]:
        """Read plain text file"""
        try:
            full_path = self._resolve_path(path)
            
            with open(full_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            # Simple section detection (by blank lines)
            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
            
            sections = []
            for i, para in enumerate(paragraphs[:50], 1):  # Limit sections
                sections.append(DocumentSection(
                    page=1,
                    heading="" if i > 1 else "Header",
                    text=para,
                    section_type="paragraph",
                    metadata={}
                ))
            
            return {
                "doc_id": path,
                "type": "text",
                "path": path,
                "sections": [
                    {
                        "page": s.page,
                        "heading": s.heading,
                        "text": s.text,
                        "type": s.section_type
                    }
                    for s in sections
                ],
                "metadata": {
                    "paragraph_count": len(paragraphs),
                    "char_count": len(text)
                },
                "extracted_findings": []
            }
            
        except Exception as e:
            return self._create_error_document(path, "text", str(e))
    
    def _resolve_path(self, path: str) -> str:
        """Resolve relative path to full path"""
        if os.path.isabs(path):
            return path
        return os.path.join(self.file_base, path)
    
    def _html_to_text(self, html: str) -> str:
        """Simple HTML to text conversion"""
        # Remove tags
        text = re.sub(r'<[^>]+>', '', html)
        # Decode entities
        text = text.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>')
        return text.strip()
    
    def _create_empty_document(self, path: str, ext: str) -> Dict[str, Any]:
        return {
            "doc_id": path,
            "type": ext.lstrip('.'),
            "path": path,
            "sections": [],
            "metadata": {"error": f"Unsupported file type: {ext}"},
            "extracted_findings": []
        }
    
    def _create_error_document(self, path: str, doc_type: str, error: str) -> Dict[str, Any]:
        return {
            "doc_id": path,
            "type": doc_type,
            "path": path,
            "sections": [],
            "metadata": {"error": error},
            "extracted_findings": [],
            "error": error
        }


class AnalysisAgent:
    """Analyzes documents and extracts structured information"""
    
    def __init__(self, ollama_base: str, model: str):
        self.ollama_base = ollama_base
        self.model = model
        self.normalizer = DocumentNormalizer()
    
    async def run_streaming(
        self,
        hits: List[Dict[str, Any]],
        strategy: Dict[str, Any],
    ) -> AsyncGenerator[Dict[str, Any], None]:
        """
        Analyze documents with progress streaming.
        Yields: {"type": "document_complete", "path": "...", "document": {...}}
                {"type": "extraction_complete", "documents": [...]}
        """
        analyzed = []
        
        for i, hit in enumerate(hits[:10], 1):  # Max 10 documents
            path = hit.get("path", "")
            if not path:
                continue
            
            # Normalize document
            document = self.normalizer.normalize(path)
            
            # Check if this is a FAT/SAT/TIB document that needs special extraction
            intent = strategy.get("intent", "fact_lookup")
            if intent in ["analysis", "comparison"]:
                # Deep analysis with LLM
                document = await self._extract_structured_data(document, strategy)
            else:
                # Basic extraction
                document = self._basic_extraction(document)
            
            analyzed.append(document)
            
            yield {
                "type": "document_complete",
                "path": path,
                "document": document,
                "progress": f"{i}/{min(len(hits), 10)}"
            }
        
        yield {
            "type": "extraction_complete",
            "documents": analyzed
        }
    
    def _basic_extraction(self, document: Dict[str, Any]) -> Dict[str, Any]:
        """Basic information extraction without LLM"""
        sections = document.get("sections", [])
        
        # Extract key sections (headings, first paragraph)
        key_content = []
        for s in sections[:5]:
            if s.get("type") in ["heading", "paragraph"]:
                key_content.append(s.get("text", "")[:500])
        
        document["extracted_findings"] = [{
            "type": "summary",
            "content": "\n".join(key_content)[:1000],
            "source": "basic_extraction"
        }]
        
        return document
    
    async def _extract_structured_data(
        self,
        document: Dict[str, Any],
        strategy: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Use LLM to extract structured data from document"""
        # Get document text
        full_text = document.get("metadata", {}).get("full_text", "")
        if not full_text:
            sections = document.get("sections", [])
            full_text = "\n".join(s.get("text", "") for s in sections[:20])
        
        # Truncate for LLM
        max_chars = 8000
        if len(full_text) > max_chars:
            full_text = full_text[:max_chars] + "\n... [truncated]"
        
        # Build extraction prompt based on strategy
        intent = strategy.get("intent", "fact_lookup")
        
        if "fat" in full_text.lower() or "sat" in full_text.lower() or "tib" in full_text.lower():
            extraction_prompt = self._build_fat_sat_prompt(full_text)
        else:
            extraction_prompt = self._build_generic_prompt(full_text, intent)
        
        try:
            messages = [
                {"role": "system", "content": "Du bist ein Dokumenten-Analyse-Agent. Extrahiere strukturierte Informationen. Antworte nur mit JSON."},
                {"role": "user", "content": extraction_prompt}
            ]
            
            result = await self._call_llm(messages)
            
            # DEBUG: Log raw LLM response
            import sys
            sys.stderr.write(f"🔍 ANALYSIS LLM raw ({len(result)} chars): {result[:200]}...\n")
            sys.stderr.flush()
            
            findings = self._parse_extraction(result)
            
            # DEBUG: Log parsed findings
            sys.stderr.write(f"📊 PARSED findings: {len(findings)} items\n")
            for f in findings[:3]:
                sys.stderr.write(f"   - {f.get('type')}: {str(f.get('content', f.get('description', '')))[:50]}...\n")
            sys.stderr.flush()
            
            document["extracted_findings"] = findings
            
        except Exception as e:
            import sys
            sys.stderr.write(f"❌ EXTRACTION ERROR: {e}\n")
            sys.stderr.flush()
            document["extracted_findings"] = [{
                "type": "error",
                "content": str(e),
                "source": "llm_extraction"
            }]
        
        return document
    
    def _build_fat_sat_prompt(self, text: str) -> str:
        return f"""Analysiere folgendes Test-Dokument (FAT/SAT/TIB) und extrahiere strukturierte Befunde:

DOKUMENT:
{text}

Extrahiere folgende Informationen als JSON:
{{
    "document_type": "FAT|SAT|TIB|Unknown",
    "customer": "Kundenname falls erwähnt",
    "test_date": "Datum falls vorhanden",
    "findings": [
        {{
            "category": "A|B|C|Error|Warning|Info",
            "severity": "high|medium|low",
            "description": "Beschreibung des Befunds",
            "status": "open|closed|in_progress"
        }}
    ],
    "summary": "Kurze Zusammenfassung der Testergebnisse"
}}

Regeln:
- A-Befunde = kritische Fehler
- B-Befunde = moderate Probleme  
- C-Befunde = Hinweise
- Ordne Befunde basierend auf Kontext der richtigen Kategorie zu
- Keine Erfindungen - nur was im Text steht"""
    
    def _build_generic_prompt(self, text: str, intent: str) -> str:
        return f"""Analysiere folgendes Dokument und extrahiere relevante Informationen:

DOKUMENT:
{text}

Extrahiere als JSON:
{{
    "key_facts": ["Fakt 1", "Fakt 2"],
    "entities": ["Entität 1", "Entität 2"],
    "summary": "Zusammenfassung",
    "relevant_quotes": ["Wichtiges Zitat 1", "Zitat 2"]
}}

Intent: {intent}"""
    
    async def _call_llm(self, messages: List[Dict[str, str]]) -> str:
        """Call Ollama"""
        payload = {
            "model": self.model,
            "messages": messages,
            "stream": False,
            "options": {"temperature": 0.2, "num_predict": 2048}
        }
        
        import httpx
        async with httpx.AsyncClient(timeout=httpx.Timeout(120.0, connect=10.0)) as client:
            r = await client.post(f"{self.ollama_base}/api/chat", json=payload)
            r.raise_for_status()
            data = r.json()
            return data.get("message", {}).get("content", "")
    
    def _parse_extraction(self, response: str) -> List[Dict[str, Any]]:
        """Parse extraction JSON"""
        try:
            data = json.loads(response.strip())
            
            # Convert to findings format
            findings = []
            
            if "findings" in data:
                for f in data["findings"]:
                    findings.append({
                        "type": "finding",
                        "category": f.get("category", "unknown"),
                        "severity": f.get("severity", "unknown"),
                        "description": f.get("description", ""),
                        "status": f.get("status", "unknown"),
                        "source": "document_analysis"
                    })
            
            if "key_facts" in data:
                for f in data["key_facts"]:
                    findings.append({
                        "type": "fact",
                        "content": f,
                        "source": "document_analysis"
                    })
            
            if "summary" in data:
                findings.append({
                    "type": "summary",
                    "content": data["summary"],
                    "source": "document_analysis"
                })
            
            return findings
            
        except json.JSONDecodeError:
            # Return raw as fallback
            return [{
                "type": "raw_extraction",
                "content": response[:1000],
                "source": "llm_raw"
            }]
